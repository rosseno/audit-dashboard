import streamlit as st
import pandas as pd
import plotly.express as px
import os
import io
from docx import Document

# Konfigurasi Halaman
st.set_page_config(page_title="Executive Audit Dashboard SPI", layout="wide")

ADMIN_PASSWORD = "SPI2026"
EXCEL_FILE = "Master_Database_Temuan_Audit_2024_2025_PSM_Ringkas.xlsx"
GDRIVE_FOLDER_DEFAULT = "https://drive.google.com/drive/folders/contoh-folder-spi-anda"

# Load Data
@st.cache_data
def load_data():
    if os.path.exists(EXCEL_FILE):
        return pd.read_excel(EXCEL_FILE)
    return pd.DataFrame()

df_master = load_data()

col_status = "Status" if "Status" in df_master.columns else (df_master.columns[-1] if not df_master.empty else "Status")
col_bidang = "Bidang" if "Bidang" in df_master.columns else (df_master.columns[5] if len(df_master.columns) > 5 else "Bidang")
col_periode = "Tahun Audit" if "Tahun Audit" in df_master.columns else (df_master.columns[3] if len(df_master.columns) > 3 else "Tahun Audit")
col_rekomendasi = "Rekomendasi Utama / Tindak Lanjut" if "Rekomendasi Utama / Tindak Lanjut" in df_master.columns else df_master.columns[4]

# Inisialisasi Session State
if "admin_auth" not in st.session_state:
    st.session_state.admin_auth = False
if "active_filter" not in st.session_state:
    st.session_state.active_filter = "ALL"

# --- SIDEBAR KONTROL ---
st.sidebar.title("🛡️ AUDIT CONTROL")
st.sidebar.markdown("---")

role = st.sidebar.selectbox("Pilih Peran / Jabatan:", [
    "Admin SPI", 
    "Direktur Utama", 
    "Direktur Operasi & Komersial", 
    "Direktur Keuangan, SDM, dll", 
    "Auditee"
])

# Autentikasi Admin di Sidebar
if role == "Admin SPI":
    if not st.session_state.admin_auth:
        pw_input = st.sidebar.text_input("Masukkan Password Admin:", type="password")
        if st.sidebar.button("Login Admin"):
            if pw_input == ADMIN_PASSWORD:
                st.session_state.admin_auth = True
                st.sidebar.success("Login Berhasil!")
                st.rerun()
            else:
                st.sidebar.error("Password Salah!")
    else:
        st.sidebar.success("Status: Admin Logged In")
        if st.sidebar.button("Logout"):
            st.session_state.admin_auth = False
            st.rerun()

is_admin_logged = (role == "Admin SPI" and st.session_state.admin_auth)

# Filter Periode
periode_options = ["Semua"]
if not df_master.empty and col_periode in df_master.columns:
    periode_options += sorted(df_master[col_periode].dropna().astype(str).unique().tolist())
selected_periode = st.sidebar.selectbox("Periode Tahun Audit:", periode_options)

# Filter Unit berdasarkan Periode
dff_filter = df_master.copy()
if selected_periode != "Semua":
    dff_filter = dff_filter[dff_filter[col_periode].astype(str) == str(selected_periode)]
bidang_options = ["Semua"] + sorted(dff_filter[col_bidang].dropna().astype(str).unique().tolist()) if not dff_filter.empty else ["Semua"]
selected_bidang = st.sidebar.selectbox("Pilih Unit:", bidang_options)

st.sidebar.markdown("---")
menu = st.sidebar.selectbox("Pilih Menu Utama:", ["Dashboard Temuan", "Upload Dokumen", "Vault KKA", "LHA Generator"])

# Batasi akses menu khusus admin
if role == "Auditee" and menu in ["Vault KKA", "LHA Generator"]:
    menu = "Upload Dokumen"

if menu in ["Vault KKA", "LHA Generator"] and not is_admin_logged:
    st.error("⚠️ Akses Terbatas. Harap login sebagai Admin SPI terlebih dahulu di sidebar.")
    st.stop()

# --- FILTER UTAMA DATA ---
dff = df_master.copy()
if role == "Direktur Operasi & Komersial":
    ops_choices = ["Operasi", "Teknik", "Pemasaran"]
    dff = dff[dff[col_bidang].astype(str).str.contains('|'.join(ops_choices), case=False, na=False)]
elif role == "Direktur Keuangan, SDM, dll":
    fin_choices = ["Keuangan", "SDM", "HSSE", "IT", "PAP", "Umum", "Rumah Tangga"]
    dff = dff[dff[col_bidang].astype(str).str.contains('|'.join(fin_choices), case=False, na=False)]

if selected_periode != "Semua":
    dff = dff[dff[col_periode].astype(str) == str(selected_periode)]
if selected_bidang != "Semua":
    dff = dff[dff[col_bidang].astype(str) == str(selected_bidang)]

total_t = len(dff)
sls_t = len(dff[dff[col_status].astype(str).str.contains("Selesai|SLS", case=False, na=False)]) if not dff.empty else 0
eval_t = len(dff[dff[col_status].astype(str).str.contains("Evaluasi|EVAL", case=False, na=False)]) if not dff.empty else 0
bd_count = len(dff[dff[col_status].astype(str).str.contains("BD|Belum|Overdue", case=False, na=False)]) if not dff.empty else 0

# --- BANNER & KOTAK METRIK ---
st.markdown(f"### 🛡️ SMART AUDIT MONITORING")
st.info(f"**Role:** {role} | **Periode:** {selected_periode} | **Unit:** {selected_bidang}")

col1, col2, col3, col4 = st.columns(4)
with col1:
    if st.button("📊 TOTAL TEMUAN\n\n**" + str(total_t) + "**", use_container_width=True):
        st.session_state.active_filter = "ALL"
with col2:
    if st.button("✅ SELESAI (SLS)\n\n**" + str(sls_t) + "**", use_container_width=True):
        st.session_state.active_filter = "SLS"
with col3:
    if st.button("⏳ EVALUASI\n\n**" + str(eval_t) + "**", use_container_width=True):
        st.session_state.active_filter = "EVAL"
with col4:
    if st.button("🚨 OVERDUE / BD\n\n**" + str(bd_count) + "**", use_container_width=True):
        st.session_state.active_filter = "BD"

st.markdown("---")

# --- RENDER KONTEN BERDASARKAN MENU ---
if menu == "Dashboard Temuan":
    if bd_count > 0:
        st.warning(f"🚨 **PERINGATAN:** ADA {bd_count} REKOMENDASI OVERDUE (BELUM DITINDAKLANJUTI)")

    col_g1, col_g2 = st.columns(2)
    with col_g1:
        if selected_bidang != "Semua" and not dff.empty:
            status_counts = dff[col_status].value_counts().reset_index()
            status_counts.columns = ['Status', 'Count']
            fig_chart = px.bar(status_counts, x='Status', y='Count', color='Status', title=f"Komposisi Status - {selected_bidang}", template="plotly_dark")
        else:
            fig_chart = px.bar(dff, x=col_bidang, color=col_status, title="Temuan per Unit", template="plotly_dark") if not dff.empty else px.bar(title="Data Kosong")
        st.plotly_chart(fig_chart, use_container_width=True)

    with col_g2:
        if not dff.empty:
            fig_pie = px.pie(dff, names=col_status, hole=0.4, title="Proporsi Status Temuan", template="plotly_dark")
            st.plotly_chart(fig_pie, use_container_width=True)
        else:
            st.info("Data proporsi tidak tersedia.")

    df_table_filtered = dff.copy()
    if st.session_state.active_filter == "SLS":
        df_table_filtered = dff[dff[col_status].astype(str).str.contains("Selesai|SLS", case=False, na=False)]
    elif st.session_state.active_filter == "EVAL":
        df_table_filtered = dff[dff[col_status].astype(str).str.contains("Evaluasi|EVAL", case=False, na=False)]
    elif st.session_state.active_filter == "BD":
        df_table_filtered = dff[dff[col_status].astype(str).str.contains("BD|Belum|Overdue", case=False, na=False)]

    st.subheader(f"📋 Rincian Database Temuan Audit (Ditampilkan: {len(df_table_filtered)} Baris)")
    
    table_cols = [col_periode, "Nama Entitas", col_bidang, col_rekomendasi, col_status]
    existing_cols = [c for c in table_cols if c in df_table_filtered.columns]
    df_display = df_table_filtered[existing_cols].copy()
    df_display.insert(0, "No", range(1, len(df_display) + 1))

    search_q = st.text_input("🔍 Cari Temuan / Kondisi / Unit:")
    if search_q:
        mask = df_display.astype(str).apply(lambda x: x.str.contains(search_q, case=False, na=False)).any(axis=1)
        df_display = df_display[mask]

    st.dataframe(df_display, use_container_width=True, height=600, hide_index=True)

elif menu == "Upload Dokumen":
    st.subheader("📎 Integrasi Bukti Tindak Lanjut ke Google Drive")
    st.markdown(f"Auditee dari unit **{selected_bidang}** dapat melampirkan tautan Google Drive berisi dokumen bukti tindak lanjut temuan audit.")
    
    # FILTER OTOMATIS: Rekomendasi temuan disaring sesuai unit/bidang yang dipilih di sidebar
    if selected_bidang != "Semua":
        dff_upload = df_master[df_master[col_bidang].astype(str) == str(selected_bidang)]
    else:
        dff_upload = df_master.copy()
        
    rekomendasi_list = dff_upload[col_rekomendasi].dropna().unique().tolist() if not dff_upload.empty else []
    
    if rekomendasi_list:
        selected_rec = st.selectbox("Pilih Rekomendasi Temuan:", rekomendasi_list)
    else:
        st.warning(f"Tidak ada rekomendasi temuan untuk unit '{selected_bidang}'.")
        selected_rec = None
        
    gdrive_link = st.text_input("🔗 Masukkan Link Google Drive (Folder/File Bukti Tindak Lanjut):", placeholder="https://drive.google.com/...")
    
    if st.button("Simpan Tautan Google Drive"):
        if selected_rec and gdrive_link:
            st.success(f"✅ Tautan Google Drive berhasil disimpan untuk temuan tersebut!")
            st.markdown(f"👉 [Buka Dokumen di Google Drive]({gdrive_link})", unsafe_allow_html=True)
        else:
            st.error("⚠️ Harap pilih rekomendasi temuan dan masukkan tautan Google Drive dengan benar.")

elif menu == "Vault KKA":
    st.subheader("📋 Vault Penyimpanan KKA & Program Audit (Terhubung Google Drive)")
    kka_link = st.text_input("🔗 Masukkan Tautan Folder Google Drive KKA / AP:", value=GDRIVE_FOLDER_DEFAULT)
    st.markdown(f"📂 **Akses Cepat Arsip KKA/AP:** [Buka Folder Google Drive SPI]({kka_link})", unsafe_allow_html=True)

elif menu == "LHA Generator":
    st.subheader("📁 LHA Generator Professional & Arsip Google Drive")
    
    p1 = st.text_area("1. Paragraf Pengantar:", placeholder="Ketik kalimat pengantar...")
    up_t1 = st.file_uploader("Upload Excel Tabel 1 untuk LHA", type=["xlsx", "xls"], key="t1")
    df_t1 = pd.read_excel(up_t1) if up_t1 else None
    if df_t1 is not None:
        st.success(f"Tabel 1 dimuat ({len(df_t1)} baris)")
        st.dataframe(df_t1.head(3))

    p2 = st.text_area("2. Uraian / Penjelasan Lanjutan:", placeholder="Ketik uraian tambahan...")
    up_t2 = st.file_uploader("Upload Excel Tabel 2 untuk LHA (Opsional)", type=["xlsx", "xls"], key="t2")
    df_t2 = pd.read_excel(up_t2) if up_t2 else None
    if df_t2 is not None:
        st.success(f"Tabel 2 dimuat ({len(df_t2)} baris)")
        st.dataframe(df_t2.head(3))

    p3 = st.text_area("3. Paragraf Kesimpulan / Penutup:", placeholder="Ketik kesimpulan...")

    lha_gdrive = st.text_input("🔗 Tautan Folder Google Drive Penyimpanan Arsip LHA:", value=GDRIVE_FOLDER_DEFAULT)
    st.markdown(f"📂 **Arsip LHA:** [Buka Folder LHA di Google Drive]({lha_gdrive})", unsafe_allow_html=True)

    if st.button("💾 Simpan & Download LHA Lengkap (.docx)", use_container_width=True):
        doc = Document()
        doc.add_heading('LAPORAN HASIL AUDIT (LHA)', 0)
        
        if p1: doc.add_paragraph(p1)
        if df_t1 is not None:
            tbl1 = doc.add_table(rows=1, cols=len(df_t1.columns))
            for i, col in enumerate(df_t1.columns): tbl1.rows[0].cells[i].text = str(col)
            for _, row in df_t1.iterrows():
                rcells = tbl1.add_row().cells
                for i, col in enumerate(df_t1.columns): rcells[i].text = str(row[col])
                
        if p2: doc.add_paragraph(p2)
        if df_t2 is not None:
            tbl2 = doc.add_table(rows=1, cols=len(df_t2.columns))
            for i, col in enumerate(df_t2.columns): tbl2.rows[0].cells[i].text = str(col)
            for _, row in df_t2.iterrows():
                rcells = tbl2.add_row().cells
                for i, col in enumerate(df_t2.columns): rcells[i].text = str(row[col])
                
        if p3: doc.add_paragraph(p3)
        
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        st.download_button(
            label="Klik untuk Unduh File .docx",
            data=stream.getvalue(),
            file_name="LHA_Multi_Tabel_Professional.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
