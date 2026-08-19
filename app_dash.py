import dash
ADMIN_PASSWORD = "SPI2026"
from dash import dcc, html, dash_table
from dash.dependencies import Input, Output, State
import plotly.express as px
import pandas as pd
import os
import base64
import io
from docx import Document

app = dash.Dash(__name__, external_stylesheets=["https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css"])
server = app.server
UPLOAD_DIRECTORY = "app_uploads"
if not os.path.exists(UPLOAD_DIRECTORY):
    os.makedirs(UPLOAD_DIRECTORY)

app.index_string = '''
<!DOCTYPE html>
<html>
    <head>
        <title>Executive Audit Dashboard</title>
        {%metas%}
        {%favicon%}
        {%css%}
        <style>
            @keyframes blink-animation {
                0% { opacity: 1; border-color: #ef4444; box-shadow: 0 0 15px rgba(239, 68, 68, 0.6); }
                50% { opacity: 0.4; border-color: #7f1d1d; box-shadow: 0 0 2px rgba(239, 68, 68, 0.1); }
                100% { opacity: 1; border-color: #ef4444; box-shadow: 0 0 15px rgba(239, 68, 68, 0.6); }
            }
            .alert-blink {
                background: rgba(239, 68, 68, 0.15);
                border: 2px solid #ef4444;
                padding: 15px 20px;
                border-radius: 8px;
                margin-bottom: 20px;
                display: flex;
                align-items: center;
                gap: 15px;
                animation: blink-animation 1.5s infinite ease-in-out;
            }
            .metric-card-btn {
                background-color: #1e293b;
                border: 1px solid #334155;
                border-radius: 8px;
                padding: 15px;
                text-align: left;
                width: 100%;
                cursor: pointer;
                transition: all 0.2s ease-in-out;
            }
            .metric-card-btn:hover {
                background-color: #334155;
                border-color: #3b82f6;
                transform: translateY(-2px);
                box-shadow: 0 4px 12px rgba(59, 130, 246, 0.3);
            }
            .upload-box {
                border: 2px dashed #475569;
                padding: 30px;
                text-align: center;
                color: #94a3b8;
                cursor: pointer;
                border-radius: 8px;
                background-color: #1e293b;
                margin-top: 15px;
            }
            .upload-box:hover {
                border-color: #3b82f6;
                color: #f8fafc;
            }
        </style>
    </head>
    <body>
        {%app_entry%}
        <footer>
            {%config%}
            {%scripts%}
            {%renderer%}
        </footer>
    </body>
</html>
'''

EXCEL_FILE = "Master_Database_Temuan_Audit_2024_2025_PSM_Ringkas.xlsx"

def load_data():
    if os.path.exists(EXCEL_FILE):
        return pd.read_excel(EXCEL_FILE)
    return pd.DataFrame()

df_master = load_data()
col_status = "Status" if "Status" in df_master.columns else (df_master.columns[-1] if not df_master.empty else "Status")
col_bidang = "Bidang" if "Bidang" in df_master.columns else (df_master.columns[5] if len(df_master.columns) > 5 else "Bidang")
col_periode = "Tahun Audit" if "Tahun Audit" in df_master.columns else (df_master.columns[3] if len(df_master.columns) > 3 else "Tahun Audit")

periode_options = [{"label": "Semua Periode", "value": "Semua"}]
if not df_master.empty and col_periode in df_master.columns:
    for p in sorted(df_master[col_periode].dropna().astype(str).unique()):
        periode_options.append({"label": p, "value": p})

app.layout = html.Div(style={"backgroundColor": "#0f172a", "minHeight": "100vh", "color": "#f1f5f9", "padding": "20px"}, children=[
    dcc.Store(id="active-filter", data="ALL"),
    dcc.Store(id="stored-table-1", data=None),
    dcc.Store(id="stored-table-2", data=None),
    dcc.Store(id="admin-auth", data=False),

    html.Div(style={"width": "280px", "float": "left", "backgroundColor": "#1e293b", "padding": "20px", "borderRadius": "10px"}, children=[
        html.H4("🛡️ AUDIT CONTROL", style={"color": "#e2e8f0"}),
        html.Hr(),
        html.Label("Pilih Peran / Jabatan:"),
        dcc.Dropdown(
            id="role",
            options=[
                {"label": "Admin SPI", "value": "admin"},
                {"label": "Direktur Utama", "value": "dirut"},
                {"label": "Direktur Operasi & Komersial", "value": "dirops"},
                {"label": "Direktur Keuangan, SDM, dll", "value": "dirkeu"},
                {"label": "Auditee", "value": "auditee"}
            ],
            value="admin",
            style={"color": "#000", "marginBottom": "15px"}
        ),
        html.Div(id="login-container", style={"marginTop": "20px"}),
            
        html.Label("Periode Tahun Audit:"),
        dcc.Dropdown(
            id="filter-periode",
            options=periode_options,
            value="Semua",
            style={"color": "#000", "marginBottom": "15px"}
        ),

        html.Label("Pilih Unit:"),
        dcc.Dropdown(
            id="filter-bidang",
            options=[{"label": "Semua Unit", "value": "Semua"}],
            value="Semua",
            style={"color": "#000", "marginBottom": "15px"}
        ),

        html.Hr(),
        html.Label("Pilih Menu Utama:"),
        dcc.Dropdown(
            id="menu",
            options=[
                {"label": "Dashboard Temuan", "value": "dash"},
                {"label": "Upload Dokumen", "value": "upload"},
                {"label": "Vault KKA", "value": "kka"},
                {"label": "LHA Generator", "value": "lha"}
            ],
            value="dash",
            style={"color": "#000", "marginBottom": "15px"}
        ),
    ]),
    html.Div(style={"marginLeft": "310px"}, children=[
        html.Div(id="role-banner", style={"background": "#1e40af", "padding": "20px", "borderRadius": "10px", "marginBottom": "20px"}),
        
        html.Div(className="row g-3 mb-4", children=[
            html.Div(className="col-md-3", children=[
                html.Button(id="btn-total", className="metric-card-btn", children=[
                    html.H6("TOTAL TEMUAN", style={"color":"#94a3b8", "fontSize":"12px", "marginBottom":"5px"}),
                    html.H3(id="val-total", children="0", style={"color":"#fff", "margin":"0"})
                ])
            ]),
            html.Div(className="col-md-3", children=[
                html.Button(id="btn-sls", className="metric-card-btn", children=[
                    html.H6("SELESAI (SLS)", style={"color":"#00CC96", "fontSize":"12px", "marginBottom":"5px"}),
                    html.H3(id="val-sls", children="0", style={"color":"#00CC96", "margin":"0"})
                ])
            ]),
            html.Div(className="col-md-3", children=[
                html.Button(id="btn-eval", className="metric-card-btn", children=[
                    html.H6("EVALUASI", style={"color":"#FFA15A", "fontSize":"12px", "marginBottom":"5px"}),
                    html.H3(id="val-eval", children="0", style={"color":"#FFA15A", "margin":"0"})
                ])
            ]),
            html.Div(className="col-md-3", children=[
                html.Button(id="btn-bd", className="metric-card-btn", children=[
                    html.H6("OVERDUE / BD", style={"color":"#EF553B", "fontSize":"12px", "marginBottom":"5px"}),
                    html.H3(id="val-bd", children="0", style={"color":"#EF553B", "margin":"0"})
                ])
            ]),
        ]),

        html.Div(id="upload-container-wrapper", style={"display": "none"}, children=[
            html.Div(style={"background": "#1e293b", "padding": "25px", "borderRadius": "8px", "border": "1px solid #334155"}, children=[
                html.H3("📎 Upload Dokumen Bukti Tindak Lanjut", style={"fontSize": "18px", "marginBottom": "10px"}),
                html.Label("Pilih Rekomendasi Temuan:"),
                dcc.Dropdown(
                    id="input-id-temuan",
                    options=[{"label": rec, "value": rec} for rec in df_master["Rekomendasi Utama / Tindak Lanjut"].dropna().unique()],
                    placeholder="Pilih rekomendasi yang akan ditindaklanjuti...",
                    style={"width": "100%", "marginBottom": "15px", "color": "black"}
                ),
                html.P("Unggah file laporan atau bukti dokumen penindaklanjutan temuan audit ke sistem.", style={"color": "#94a3b8", "fontSize": "13px"}),
                dcc.Upload(
                    id="upload-data",
                    children=html.Div([
                        html.B("Seret & Letakkan file di sini"), html.Span(" atau "), html.U("Pilih File")
                    ]),
                    className="upload-box",
                    multiple=False
                ),
                html.Div(id="upload-status")
            ])
        ]),

        html.Div(id="main-content")
    ])
])

@app.callback(
    Output("active-filter", "data"),
    [
        Input("btn-total", "n_clicks"),
        Input("btn-sls", "n_clicks"),
        Input("btn-eval", "n_clicks"),
        Input("btn-bd", "n_clicks")
    ],
    [
        State("active-filter", "data")
    ]
)
def update_filter(n_total, n_sls, n_eval, n_bd, current_filter):
    ctx = dash.callback_context
    if not ctx.triggered:
        return "ALL"
    button_id = ctx.triggered[0]['prop_id'].split('.')[0]
    if button_id == "btn-total":
        return "ALL"
    elif button_id == "btn-sls":
        return "SLS"
    elif button_id == "btn-eval":
        return "EVAL"
    elif button_id == "btn-bd":
        return "BD"
    return "ALL"

@app.callback(
    [Output("filter-bidang", "options"), Output("filter-bidang", "value")],
    [Input("filter-periode", "value")]
)
def update_bidang_dropdown(selected_periode):
    if df_master.empty:
        return [{"label": "Semua Unit", "value": "Semua"}], "Semua"
    dff = df_master.copy()
    if selected_periode and selected_periode != "Semua":
        dff = dff[dff[col_periode].astype(str) == str(selected_periode)]
    bidang_list = sorted(dff[col_bidang].dropna().astype(str).unique())
    options = [{"label": "Semua Unit", "value": "Semua"}] + [{"label": b, "value": b} for b in bidang_list]
    return options, "Semua"

@app.callback(
    [Output("main-content", "children"),
     Output("upload-container-wrapper", "style"),
     Output("role-banner", "children"),
     Output("val-total", "children"),
     Output("val-sls", "children"),
     Output("val-eval", "children"),
     Output("val-bd", "children")],
    [Input("menu", "value"), Input("role", "value"), Input("filter-periode", "value"), 
     Input("filter-bidang", "value"), Input("active-filter", "data"), Input("admin-auth", "data")]
)
def update(menu, role, selected_periode, selected_bidang, active_filter, is_auth):
    is_admin_logged = (role == "admin" and is_auth)
    
    if role == "auditee" and menu in ["kka", "lha"]:
        menu = "upload"
    
    if menu in ["kka", "lha"] and not is_admin_logged:
        return [html.Div("⚠️ Akses Terbatas. Harap masukkan password Admin terlebih dahulu.", style={"color":"#ef4444", "padding":"20px"})], {"display":"none"}, [html.H2("Akses Ditolak")], "0", "0", "0", "0"

    if df_master.empty:
        return html.Div([html.P("File database excel tidak ditemukan / kosong.")]), {"display": "none"}, [html.H2("SMART AUDIT MONITORING")], "0", "0", "0", "0"

    dff = df_master.copy()
    
    if role == "dirops":
        ops_choices = ["Operasi", "Teknik", "Pemasaran"]
        dff = dff[dff[col_bidang].astype(str).str.contains('|'.join(ops_choices), case=False, na=False)]
    elif role == "dirkeu":
        fin_choices = ["Keuangan", "SDM", "HSSE", "IT", "PAP", "Umum", "Rumah Tangga"]
        dff = dff[dff[col_bidang].astype(str).str.contains('|'.join(fin_choices), case=False, na=False)]

    if selected_periode and selected_periode != "Semua":
        dff = dff[dff[col_periode].astype(str) == str(selected_periode)]

    total_t = len(dff)
    sls_t = len(dff[dff[col_status].astype(str).str.contains("Selesai|SLS", case=False, na=False)]) if not dff.empty else 0
    eval_t = len(dff[dff[col_status].astype(str).str.contains("Evaluasi|EVAL", case=False, na=False)]) if not dff.empty else 0
    bd_count = len(dff[dff[col_status].astype(str).str.contains("BD|Belum|Overdue", case=False, na=False)]) if not dff.empty else 0
    
    banner = [html.H2("SMART AUDIT MONITORING"), html.Div(f"Role: {(role or 'admin').upper()} | Periode: {selected_periode} | Unit: {selected_bidang} | Filter: {active_filter}")]
    
    upload_style = {"display": "none"}
    content = html.Div()

    if menu == "dash":
        alert_box = html.Div(className="alert-blink", children=[
            html.Span("🚨", style={"fontSize": "24px"}),
            html.Div([
                html.Div(f"PERINGATAN: ADA {bd_count} REKOMENDASI OVERDUE (BELUM DITINDAKLANJUTI)", style={"color": "#f87171", "fontWeight": "700", "fontSize": "15px"})
            ])
        ]) if bd_count > 0 else html.Div()

        df_table_filtered = dff.copy()
        if active_filter == "SLS":
            df_table_filtered = dff[dff[col_status].astype(str).str.contains("Selesai|SLS", case=False, na=False)]
        elif active_filter == "EVAL":
            df_table_filtered = dff[dff[col_status].astype(str).str.contains("Evaluasi|EVAL", case=False, na=False)]
        elif active_filter == "BD":
            df_table_filtered = dff[dff[col_status].astype(str).str.contains("BD|Belum|Overdue", case=False, na=False)]

        table_cols = [col_periode, "Nama Entitas", col_bidang, "Rekomendasi Utama / Tindak Lanjut", col_status]
        existing_cols = [c for c in table_cols if c in df_table_filtered.columns]
        
        df_display = df_table_filtered[existing_cols].copy()
        records = df_display.to_dict('records')
        for idx, row in enumerate(records, start=1):
            row["No"] = idx

        table_columns = [{"name": "No", "id": "No"}] + [{"name": c, "id": c} for c in existing_cols]

        if selected_bidang != "Semua" and not dff.empty:
            status_counts = dff[col_status].value_counts().reset_index()
            status_counts.columns = ['Status', 'Count']
            fig_chart = px.bar(status_counts, x='Status', y='Count', color='Status', title=f"Komposisi Status Temuan - {selected_bidang}", template="plotly_dark")
            fig_chart.update_layout(xaxis_title=None)
        else:
            fig_chart = px.bar(dff, x=col_bidang, color=col_status, title="Temuan per Unit", template="plotly_dark") if not dff.empty else {}
            if not df_master.empty:
                fig_chart.update_layout(xaxis_title=None)

        content = [
            alert_box,
            dcc.Graph(figure=fig_chart),
            html.Br(),
            html.Div(style={"background": "#1e293b", "padding": "20px", "borderRadius": "8px", "border": "1px solid #334155"}, children=[
                html.H4(f"Ringkasan Data Temuan (Periode: {selected_periode})", style={"fontSize": "16px", "marginBottom": "15px"}),
                dash_table.DataTable(
                    data=records,
                    columns=table_columns,
                    page_size=10,
                    style_table={'overflowX': 'auto'},
                    style_header={"backgroundColor": "#0f172a", "color": "white", "fontWeight": "bold"},
                    style_cell={'backgroundColor': '#1e293b', 'color': 'white', 'textAlign': 'left', 'padding': '10px'}
                )
            ])
        ]
    elif menu == "upload":
        upload_style = {"display": "block"}
        content = html.Div()
    elif menu == "kka":
        content = html.Div([
            html.H3("📋 Vault Penyimpanan File KKA & Program Audit (AP)", style={"fontSize": "18px", "marginBottom": "10px"}),
            html.P("Pusat arsip dokumen kertas kerja audit dan program audit.", style={"color": "#94a3b8"})
        ])
    elif menu == "lha":
        content = html.Div([
            html.H3("📁 LHA Generator Professional (Multi-Tabel & Narasi)", style={"fontSize": "18px", "marginBottom": "20px"}),
            
            html.Label("1. Paragraf Pengantar:", style={"color": "#38bdf8"}),
            dcc.Textarea(id="lha-p1", placeholder="Ketik kalimat pengantar...", style={"width": "100%", "height": "80px", "backgroundColor": "#0f172a", "color": "white", "marginBottom": "15px"}),
            
            html.Div(style={"margin": "10px 0", "padding": "15px", "border": "1px dashed #38bdf8"}, children=[
                html.Label("Upload Tabel Pertama:", style={"fontWeight": "bold"}),
                dcc.Upload(id="upload-table-lha-1", children=html.Button("Upload Excel Tabel 1", style={"backgroundColor":"#334155", "color":"white", "border":"none", "padding":"6px 12px", "borderRadius":"5px"})),
                html.Div(id="preview-table-lha-1", style={"marginTop": "10px"})
            ]),
            
            html.Label("2. Uraian / Penjelasan Lanjutan:", style={"color": "#38bdf8", "marginTop": "15px"}),
            dcc.Textarea(id="lha-p2", placeholder="Ketik uraian tambahan di sini...", style={"width": "100%", "height": "100px", "backgroundColor": "#0f172a", "color": "white", "marginBottom": "15px"}),
            
            html.Div(style={"margin": "10px 0", "padding": "15px", "border": "1px dashed #38bdf8"}, children=[
                html.Label("Upload Tabel Berikutnya (Opsional):", style={"fontWeight": "bold"}),
                dcc.Upload(id="upload-table-lha-2", children=html.Button("Upload Excel Tabel 2", style={"backgroundColor":"#334155", "color":"white", "border":"none", "padding":"6px 12px", "borderRadius":"5px"})),
                html.Div(id="preview-table-lha-2", style={"marginTop": "10px"})
            ]),

            html.Label("3. Paragraf Kesimpulan / Penutup:", style={"color": "#38bdf8", "marginTop": "15px"}),
            dcc.Textarea(id="lha-p3", placeholder="Ketik kesimpulan...", style={"width": "100%", "height": "80px", "backgroundColor": "#0f172a", "color": "white", "marginBottom": "15px"}),
            
            html.Button("💾 Simpan & Download LHA Lengkap (.docx)", id="btn-save-lha", n_clicks=0, style={"marginTop": "10px", "backgroundColor": "#2563eb", "color": "white", "padding": "12px 20px", "width": "100%", "border": "none", "borderRadius": "5px", "fontWeight": "bold", "cursor": "pointer"}),
            dcc.Download(id="download-lha-docx")
        ])
    else:
        content = html.Div([
            html.H3("📁 Vault Penyimpanan File LHA Word", style={"fontSize": "18px", "marginBottom": "10px"}),
            html.P("Pusat arsip Laporan Hasil Audit resmi.", style={"color": "#94a3b8"})
        ])
    
    return content, upload_style, banner, str(total_t), str(sls_t), str(eval_t), str(bd_count)

@app.callback(
    Output("upload-status", "children"),
    [Input("upload-data", "contents")],
    [State("upload-data", "filename"), State("input-id-temuan", "value")],
    prevent_initial_call=True
)
def save_file(contents, filename, id_temuan):
    if not id_temuan:
        return html.Div("⚠️ Harap pilih rekomendasi temuan terlebih dahulu sebelum mengunggah file!", style={"color": "#EF553B", "marginTop": "15px", "fontWeight": "bold"})
    
    temuan_folder = os.path.join(UPLOAD_DIRECTORY, str(id_temuan))
    if not os.path.exists(temuan_folder):
        os.makedirs(temuan_folder)
        
    if contents is not None:
        content_type, content_string = contents.split(',')
        decoded = base64.b64decode(content_string)
        with open(os.path.join(temuan_folder, filename), 'wb') as f:
            f.write(decoded)
        return html.Div(f"✅ Berhasil! File {filename} tersimpan untuk rekomendasi tersebut.", style={"color": "#00CC96", "marginTop": "15px", "fontWeight": "bold"})
    return ""

@app.callback(
    [Output("preview-table-lha-1", "children"), Output("stored-table-1", "data")],
    [Input("upload-table-lha-1", "contents")], [State("upload-table-lha-1", "filename")]
)
def preview_t1(contents, filename):
    if not contents: return "", None
    df = pd.read_excel(io.BytesIO(base64.b64decode(contents.split(',')[1])))
    records = df.to_dict('records')
    cols = [{"name": c, "id": c} for c in df.columns]
    table = dash_table.DataTable(data=records, columns=cols, page_size=3, style_table={'overflowX': 'auto'}, style_header={"backgroundColor": "#0f172a", "color": "white"}, style_cell={'backgroundColor': '#1e293b', 'color': 'white'})
    return html.Div([html.P(f"📁 Terunggah: {filename}", style={"color": "#00CC96", "fontSize": "12px"}), table]), records

@app.callback(
    [Output("preview-table-lha-2", "children"), Output("stored-table-2", "data")],
    [Input("upload-table-lha-2", "contents")], [State("upload-table-lha-2", "filename")]
)
def preview_t2(contents, filename):
    if not contents: return "", None
    df = pd.read_excel(io.BytesIO(base64.b64decode(contents.split(',')[1])))
    records = df.to_dict('records')
    cols = [{"name": c, "id": c} for c in df.columns]
    table = dash_table.DataTable(data=records, columns=cols, page_size=3, style_table={'overflowX': 'auto'}, style_header={"backgroundColor": "#0f172a", "color": "white"}, style_cell={'backgroundColor': '#1e293b', 'color': 'white'})
    return html.Div([html.P(f"📁 Terunggah: {filename}", style={"color": "#00CC96", "fontSize": "12px"}), table]), records

@app.callback(
    Output("download-lha-docx", "send_bytes"),
    [Input("btn-save-lha", "n_clicks")],
    [State("lha-p1", "value"), State("lha-p2", "value"), State("lha-p3", "value"), 
     State("stored-table-1", "data"), State("stored-table-2", "data")],
    prevent_initial_call=True
)
def generate_docx_multi(n_clicks, p1, p2, p3, t1_data, t2_data):
    if not n_clicks: return None
    
    doc = Document()
    doc.add_heading('LAPORAN HASIL AUDIT (LHA)', 0)
    
    if p1: doc.add_paragraph(p1)
    
    if t1_data:
        df1 = pd.DataFrame(t1_data)
        tbl1 = doc.add_table(rows=1, cols=len(df1.columns))
        for i, col in enumerate(df1.columns): tbl1.rows[0].cells[i].text = str(col)
        for row in t1_data:
            rcells = tbl1.add_row().cells
            for i, col in enumerate(df1.columns): rcells[i].text = str(row.get(col, ''))
            
    if p2: doc.add_paragraph(p2)
    
    if t2_data:
        df2 = pd.DataFrame(t2_data)
        tbl2 = doc.add_table(rows=1, cols=len(df2.columns))
        for i, col in enumerate(df2.columns): tbl2.rows[0].cells[i].text = str(col)
        for row in t2_data:
            rcells = tbl2.add_row().cells
            for i, col in enumerate(df2.columns): rcells[i].text = str(row.get(col, ''))
            
    if p3: doc.add_paragraph(p3)
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return dcc.send_bytes(stream.getvalue(), filename="LHA_Multi_Tabel_Professional.docx")

@app.callback(
    Output("login-container", "children"),
    [Input("role", "value")]
)
def show_login(role):
    if role == "admin":
        return [
            dcc.Input(id="admin-pw", type="password", placeholder="Masukkan Password...", style={"width":"100%", "marginTop":"5px", "color": "black", "backgroundColor": "white"}),
            html.Button("Login", id="btn-login", className="btn btn-primary btn-sm", style={"width":"100%", "marginTop":"5px"})
        ]
    return []

@app.callback(
    Output("admin-auth", "data"),
    [Input("btn-login", "n_clicks")],
    [
        State("admin-pw", "value"), 
        State("admin-auth", "data")
    ],
    prevent_initial_call=True
)
def authenticate(n_clicks, password, current_auth):
    if n_clicks:
        if password == ADMIN_PASSWORD:
            return True
        else:
            return False
    return False


